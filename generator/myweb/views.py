from django.shortcuts import render,redirect,HttpResponse
import cohere
from django.http import JsonResponse
import os
from django.conf import settings
import docx
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login, authenticate, logout
import json
from .models import PersonalInfo, Genders, Types, MaritalStatus, EducAttain, Ranges, EstIncome, Requests, Profile
from django.db.models import Count, Q
from datetime import datetime, timedelta
from django.core.serializers import serialize
from django.core.exceptions import ValidationError
from django.core.files.base import ContentFile
from django.utils.text import slugify



@login_required(login_url='login')  
def generator(request):
    if request.method=='POST':  
        data = json.loads(request.body)
    
        p = data.get('problem')
        a = data.get('asso')
    
        print(p)
        print(a)
        
        message = 'Associated Problem: \n \n		Interventions: \n\n		Activities: \n\n		Resources Needed: \n\n		Expected Outcome: \n\n\nAssociated Problem: \n \n 	Interventions: \n\n 	Activities: \n\n 	Responsible Person: \n\n 	Resources Needed: \n\n 	Expected Outcome: \n\n make an action plan for a person with a presenting problem' + p + ' and the associated problem(s) ' + a + '. provide interventions, activities, resources need and expected outcome for EACH associated problem, follow the format given above.\n\n'
        
        reco = 'create a recommendation for a person with a presenting problem' + p + 'and associated problem(s)' + a  

  
        co = cohere.Client('ztzx0zGkAcOk5bjq1ncUIUwnNPeJE7ob66wqgr6T') 
        response = co.chat(
            message, 
            model="command", 
            temperature=0.9
        )
        answer = response.text
        
        print(answer)
        
        co = cohere.Client('ztzx0zGkAcOk5bjq1ncUIUwnNPeJE7ob66wqgr6T') 
        responses = co.chat(
            reco, 
            model="command", 
            temperature=0.9
        )
        
        recom = responses.text
        
        print(recom)
        return JsonResponse({'answer': answer, 'recom':recom})
    

    return JsonResponse({'error': 'Invalid request method'}, status=405)
def results(request):
    doc_filename = 'format2.docx'
    doc_path2 = os.path.join(settings.STATICFILES_DIRS[0], doc_filename)
    try:
        doc2 = docx.Document(doc_path2)
        content2 = '\n'.join([paragraph.text for paragraph in doc2.paragraphs])
    except Exception as e:
        content2 = f"Error reading document: {e}"
        
    username = request.user.username
    
    return render(request, 'results.html',{'content2':content2, 'username':username})

def get_age_range(bdate):
    bday_date = datetime.strptime(bdate, '%Y-%m-%d').date()
    today_date = datetime.now().date()

    age_delta = today_date - bday_date
    age_years = age_delta.days // 365

    age_ranges_db = Ranges.objects.all()
    age_ranges = {}

    for range_obj in age_ranges_db:
        age_ranges[range_obj.ranges_name] = (int(range_obj.range1), int(range_obj.range2))

    for range_name, (min_age, max_age) in age_ranges.items():
        if min_age <= age_years <= max_age:
            findRanges = Ranges.objects.get(ranges_name=range_name)
            return findRanges

    return Ranges.objects.get(ranges_name="UNKOWN")

def edit(request):
    profile = Profile.objects.get(user=request.user)
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        email = request.POST.get('email')

        request.user.username = username
        request.user.set_password(password)
        request.user.email = email
        request.user.save()

        if 'profile_picture' in request.FILES:
            profile.profile_picture = request.FILES['profile_picture']
            profile.save()

        return redirect('home')
        
def save(request):
    if request.method == 'POST':
        date = request.POST.get('currentdate')
        caseno = request.POST.get('case')
        addressinpt = request.POST.get('address')
        typesinpt = request.POST.get('type')
        fname = request.POST.get('fname')
        lname = request.POST.get('lname')
        mname = request.POST.get('mname')
        bdate = request.POST.get('bday')
        sexgender = request.POST.get('gender')
        status = request.POST.get('marital-status')
        educ_attain = request.POST.get('educ')
        income = request.POST.get('means')
        est_income = request.POST.get('estimated')
        prblm = request.POST.get('problem')
        assoc = request.POST.get('assoc')

        print(status)

        findtypes = Types.objects.get(type_name=typesinpt)
        findGndr = Genders.objects.get(gndr_name=sexgender)
        findMartial = MaritalStatus.objects.get(maritals_name=status)
        findEduc = EducAttain.objects.get(educ_name=educ_attain)

        try:
            findEst = EstIncome.objects.get(est_income_details=est_income)
            einc = findEst.est_income_name
        except EstIncome.DoesNotExist:
            einc = "Unknown"

        if mname is None:
            full_name = f"{fname} {lname}"
        else:
            full_name = f"{fname} {mname} {lname}"

        age_range = get_age_range(bdate)


        PersonalInfo.objects.create(
                user=request.user,
                current_date=date,
                case_number=caseno,
                full_name=full_name,
                birthdate=bdate,
                range=age_range,
                type=findtypes,
                address=addressinpt,
                sex_gender=findGndr,
                marital_status=findMartial,
                educational_attainment=findEduc,
                means_of_income=income,
                estimated_income=findEst,
                money=einc,
                problem=prblm,
            )
        
        new_personal_info = PersonalInfo.objects.latest('id')
        personal_info_id = new_personal_info.id

        doc_filename = 'format2.docx'
        doc_path2 = os.path.join(settings.STATICFILES_DIRS[0], doc_filename)

        username = request.user.username

        try:
            doc2 = docx.Document(doc_path2)
            content2 = '\n'.join([paragraph.text for paragraph in doc2.paragraphs])
        except Exception as e:
            content2 = f"Error reading document: {e}"
                
        message = f'You will act as an Action Plan Generator. the client is a {age_range}, {findMartial}, {findGndr}, {findtypes} with an Educational background of {findEduc}, has a source of income of {income} Worker with an income of {einc}, make an action plan for the client with a presenting problem of {prblm} and the associated problem(s) of {assoc}. follow the format and make an action plan for every associated problems and start with providing the given format right away and avoid providing other responses in the beginning aside from the format\n\n\nAssociated Problem 1:\nInterventions:\nActivities:\nResources Needed:\nExpected Outcome:\n\nAssociated Problem 2:\nInterventions:\nActivities:\nResources Needed:\nExpected Outcome:'


        co = cohere.Client('tFlLE5DvoNaYC9L3YESDDGI9vno3xX1wDqc0pw6U')
        response = co.chat(
                message,
                model="command",
                temperature=0.9
        )
        answer = response.text
       



        

        return render(request, 'results.html', {'date': date, 'caseno':caseno,'addressinpt':addressinpt,'typesinpt':typesinpt,
                                                'full_name':full_name,'bdate':bdate,'sexgender':sexgender,'status':status,'educ_attain':educ_attain,
                                                'income':income,'est_income':est_income,'prblm':prblm,'username':username,'answer':answer,
                                                'content2':content2,'prblm':prblm,'assoc':assoc,'personal_info_id':personal_info_id}) 
        
def genRecom(request):
    
    if request.method == 'POST':
        data = json.loads(request.body)  
        prblm = data.get('prblm')
        assoc = data.get('assoc')
    
        reco = f'You will act a professional recommendation bot, make a recommendation for a person with a presenting problem {prblm} and associated problem(s) {assoc}'


        co = cohere.Client('tFlLE5DvoNaYC9L3YESDDGI9vno3xX1wDqc0pw6U')
        responses = co.chat(
                reco,
                model="command",
                temperature=0.9
            )

        recom = responses.text
    
        
        return JsonResponse({'recom':recom})



@login_required(login_url='login')  
def homepage(request):

    profile, created = Profile.objects.get_or_create(user=request.user)

    total_cases = PersonalInfo.objects.count()

    new_total_cases = int(total_cases) + int(1)

    doc_filename = 'format2.docx'
    doc_path2 = os.path.join(settings.STATICFILES_DIRS[0], doc_filename)

    user = request.user
    
    displayGenders = Genders.objects.all()
    displayMaritals = MaritalStatus.objects.all()
    displayTypes = Types.objects.all()
    displayRanges = Ranges.objects.all()
    displayEduc = EducAttain.objects.all()
    displaySocials = EstIncome.objects.all()
    
    try:
        doc2 = docx.Document(doc_path2)
        content2 = '\n'.join([paragraph.text for paragraph in doc2.paragraphs])
    except Exception as e:
        content2 = f"Error reading document: {e}"

    return render(request, 'index.html', {'content2':content2,'user': user,'new_total_cases': new_total_cases,'profile':profile,
                                          'displayGenders':displayGenders, 'displayMaritals':displayMaritals, 'displayTypes':displayTypes,
                                         'displayRanges':displayRanges,'displayEduc':displayEduc,'displaySocials':displaySocials})


def adminPage(request):
    users = User.objects.filter(is_staff=False, is_superuser=False)
    requests = Requests.objects.all()
    print(users)

    displayGenders = Genders.objects.all()
    displayMaritals = MaritalStatus.objects.all()
    displayRanges = Ranges.objects.all()
    displayEduc = EducAttain.objects.all()
    displaySocials = EstIncome.objects.all()

    folders = PersonalInfo.objects.values('problem').annotate(count=Count('problem'))
    total_cases = PersonalInfo.objects.count()
    address_data = PersonalInfo.objects.values('address').annotate(count=Count('address'))
    range_data = PersonalInfo.objects.values('range').annotate(count=Count('range'))
    range_names = [Ranges.objects.get(id=item['range']).ranges_name for item in range_data]
    social_data = PersonalInfo.objects.values('estimated_income').annotate(count=Count('estimated_income'))
    social_names = [EstIncome.objects.get(id=item['estimated_income']).est_income_name for item in social_data]
    gender_data = PersonalInfo.objects.values('sex_gender').annotate(count=Count('sex_gender'))
    gndr_names = [Genders.objects.get(id=item['sex_gender']).gndr_name for item in gender_data]
    types_data = PersonalInfo.objects.values('type').annotate(count=Count('range'))
    type_names = [Types.objects.get(id=item['type']).type_name for item in types_data]

    types_with_count = PersonalInfo.objects.values('type').annotate(count=Count('type')).order_by('-count')
    highest_types = types_with_count.first()['type'] if types_with_count else None
    most_highest_type = None

    if highest_types is not None:
        types_obj = Types.objects.filter(id=highest_types).first()
        if types_obj:
            most_highest_type = types_obj.type_name


    genders_with_count = PersonalInfo.objects.values('sex_gender').annotate(count=Count('sex_gender')).order_by('-count')
    highest_gender = genders_with_count.first()['sex_gender'] if genders_with_count else None

    most_common_gender = None

    if highest_gender is not None:
        gender_obj = Genders.objects.filter(id=highest_gender).first()
        if gender_obj:
            most_common_gender = gender_obj.gndr_name

    addresses_with_count = PersonalInfo.objects.values('address').annotate(count=Count('address')).order_by('-count')
    
    highest_count = addresses_with_count.first()['count'] if addresses_with_count else 0
    most_common_addresses = [item for item in addresses_with_count if item['count'] == highest_count]

    ranges_with_count = PersonalInfo.objects.values('range').annotate(count=Count('range')).order_by('-count')
    highest_range = ranges_with_count.first()['range'] if ranges_with_count else None

    most_common_ranges = None

    if highest_range is not None:
        range_obj = Ranges.objects.filter(id=highest_range).first()
        if range_obj:
            most_common_ranges = range_obj.ranges_name

    class_with_count = PersonalInfo.objects.values('estimated_income').annotate(count=Count('estimated_income')).order_by('-count')

    highest_est = class_with_count.first()['estimated_income'] if class_with_count else None

    most_common_est = None

    if highest_est is not None:
        est_income_obj = EstIncome.objects.filter(id=highest_est).first()
        if est_income_obj:
            most_common_est = est_income_obj.est_income_name


    return render(request, 'admin.html', {'address_data': address_data,'total_cases':total_cases,'most_common_gender':most_common_gender,
                                        'range_data': range_data,'range_names':range_names,'most_common_ranges':most_common_ranges,'most_common_est':most_common_est,
                                        'social_data':social_data, 'folders':folders, 'users': users, 'displayGenders':displayGenders,
                                        'displayMaritals':displayMaritals,'displayRanges':displayRanges,'displayEduc':displayEduc,
                                        'displaySocials':displaySocials,'requests':requests,'most_common_addresses':most_common_addresses,
                                        'most_highest_type':most_highest_type,'gender_data':gender_data,'types_with_count':types_with_count,
                                        'gndr_names':gndr_names,'social_names':social_names,'type_names':type_names}) 



def landingPage(request):
    return render(request, 'landing.html') 

 
def logins(request):
    if request.method == 'POST': 
            username = request.POST.get('username')
            pass1 = request.POST.get('password')
 
            User = authenticate(request,username=username,password=pass1) 
            if User is not None:
                if User.is_superuser and User.is_staff:
                    login(request,User)
                    return redirect('admin')
                else:
                    login(request, User)
                    return redirect('home')
            else:   
                return redirect('landing')
    return render (request, 'landing.html')


def create(request):          
    if request.method=='POST':  
        uname=request.POST.get('username')
        email=request.POST.get('email')
        pass1=request.POST.get('password')
      
        if User.objects.filter(username=uname).exists() or User.objects.filter(email=email).exists():
            return redirect('admin')
        else:
            my_user=User.objects.create_user(uname,email,pass1)
            my_user.save()
            return redirect('admin')
    return render (request, 'admin.html')

def requestAcc(request):          
    if request.method == 'POST':  
        uname = request.POST.get('username')
        email = request.POST.get('email')
        pass1 = request.POST.get('password')
      
        if Requests.objects.filter(username=uname).exists() or Requests.objects.filter(email=email).exists():
            return redirect('landing')
        else:
            my_user = Requests.objects.create(username=uname, email=email, password=pass1)
            my_user.save
            return redirect('landing')



def find_filter_id(filter2):
    try:
        findfilter2 = EstIncome.objects.get(est_income_name=filter2)
        return findfilter2.id
    except EstIncome.DoesNotExist:
        pass  
    except EstIncome.MultipleObjectsReturned:
        pass  

    try:
        findfilter2 = Genders.objects.get(gndr_name=filter2)
        return findfilter2.id
    except Genders.DoesNotExist:
        pass  
    except Genders.MultipleObjectsReturned:
        pass

    try:
        findfilter2 = Types.objects.get(type_name=filter2)
        return findfilter2.id
    except Types.DoesNotExist:
        pass  
    except Types.MultipleObjectsReturned:
        pass

    try:
        findfilter2 = Ranges.objects.get(ranges_name=filter2)
        return findfilter2.id
    except Ranges.DoesNotExist:
        pass  
    except Ranges.MultipleObjectsReturned:
        pass

    try:
        findfilter2 = MaritalStatus.objects.get(maritals_name=filter2)
        return findfilter2.id
    except MaritalStatus.DoesNotExist:
        pass  
    except MaritalStatus.MultipleObjectsReturned:
        pass 

    try:
        findfilter2 = EducAttain.objects.get(educ_name=filter2)
        return findfilter2.id
    except EducAttain.DoesNotExist:
        pass  
    except EducAttain.MultipleObjectsReturned:
        pass 

    try:
        findfilter2 = EstIncome.objects.get(est_income_name=filter2)
        return findfilter2.id
    except EstIncome.DoesNotExist:
        pass  
    except EstIncome.MultipleObjectsReturned:
        pass

    return None

def filter(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        filter1 = data.get('filter1')
        filter2 = data.get('filter2')

        print(filter1)
        print(filter2)

        filter_id = find_filter_id(filter2)



        if filter1 == 'All' and filter2 == 'All':
            gender_counts = PersonalInfo.objects.values('sex_gender').annotate(count=Count('sex_gender'))
            type_counts = PersonalInfo.objects.values('type').annotate(count=Count('type'))
            filter_social = PersonalInfo.objects.values('estimated_income').annotate(count=Count('estimated_income'))
            ranges = PersonalInfo.objects.values('range').annotate(count=Count('range'))
            marital = PersonalInfo.objects.values('marital_status').annotate(count=Count('marital_status'))
            educ = PersonalInfo.objects.values('educational_attainment').annotate(count=Count('educational_attainment'))
            gndr_names = [Genders.objects.get(id=item['sex_gender']).gndr_name for item in gender_counts]
            type_names = [Types.objects.get(id=item['type']).type_name for item in type_counts]
            social_names = [EstIncome.objects.get(id=item['estimated_income']).est_income_name for item in filter_social]
            ranges_names = [Ranges.objects.get(id=item['range']).ranges_name for item in ranges]
            maritals_names = [MaritalStatus.objects.get(id=item['marital_status']).maritals_name for item in marital]
            educ_names = [EducAttain.objects.get(id=item['educational_attainment']).educ_name for item in educ]
            means = PersonalInfo.objects.values('means_of_income').annotate(count=Count('means_of_income'))
            problems = PersonalInfo.objects.values('problem').annotate(count=Count('problem'))

            print(gender_counts)
            total_count = gender_counts.aggregate(total=Count('count'))['total']

            gender_percentages = []
            for gender_count in gender_counts:
                percentage = (gender_count['count'] / total_count) * 100
                gender_percentages.append({
                    'gender': gender_count['sex_gender'],
                    'count': gender_count['count'],
                    'percentage': percentage,
                })
            
            female_count = [item['count'] for item in gender_counts]
            
            community_count = [item['count'] for item in type_counts]

            social_data_counts = [item['count'] for item in filter_social]

            range_data_counts = [item['count'] for item in ranges]

            marital_status_counts = [item['count'] for item in marital]

            means_data = [item['means_of_income'] for item in means]
            means_data_counts = [item['count'] for item in means]

            educ_data_counts = [item['count'] for item in educ]

            prob_data = [item['problem'] for item in problems]
            prob_data_counts = [item['count'] for item in problems]

            data = {
                'gndr_names':gndr_names,
                'female_count':female_count,
                'type_names':type_names,
                'community_count':community_count,
                'social_names':social_names,
                'maritals_names':maritals_names,
                'marital_status_counts':marital_status_counts,
                'social_data_counts':social_data_counts,
                'ranges_names':ranges_names,
                'range_data_counts':range_data_counts,
                'means_data':means_data,
                'means_data_counts':means_data_counts,
                'educ_names':educ_names,
                'educ_data_counts':educ_data_counts,
                'prob_data':prob_data,
                'prob_data_counts':prob_data_counts,
                'gender_percentages': gender_percentages
            }
            return JsonResponse(data)
        
        elif filter1 == 'Gender' and filter_id == filter_id:
            gender_counts_per_marital = PersonalInfo.objects.values('marital_status').annotate(
                    gender_count=Count('sex_gender', filter=Q(sex_gender=filter_id)),
                )

            gender_counts = PersonalInfo.objects.values('sex_gender').annotate(count=Count('sex_gender'))

            gender_counts_per_social= PersonalInfo.objects.values('estimated_income').annotate(
                    social_count=Count('sex_gender', filter=Q(sex_gender=filter_id)),
                )
            gender_counts_per_range= PersonalInfo.objects.values('range').annotate(
                    range_count=Count('sex_gender', filter=Q(sex_gender=filter_id)),
                )
            gender_counts_per_type= PersonalInfo.objects.values('type').annotate(
                    type_count=Count('sex_gender', filter=Q(sex_gender=filter_id)),
                )
            gender_counts_per_attainment= PersonalInfo.objects.values('educational_attainment').annotate(
                    educ_count=Count('sex_gender', filter=Q(sex_gender=filter_id)),
                )
            gender_counts_per_means= PersonalInfo.objects.values('means_of_income').annotate(
                    means_count=Count('sex_gender', filter=Q(sex_gender=filter_id)),
                )
            gender_counts_per_prob= PersonalInfo.objects.values('problem').annotate(
                    prob_counts=Count('sex_gender', filter=Q(sex_gender=filter_id)),
                )
            
            gndr_names = [Genders.objects.get(id=item['sex_gender']).gndr_name for item in gender_counts]
            type_names = [Types.objects.get(id=item['type']).type_name for item in gender_counts_per_type]
            social_names = [EstIncome.objects.get(id=item['estimated_income']).est_income_name for item in gender_counts_per_social]
            ranges_names = [Ranges.objects.get(id=item['range']).ranges_name for item in gender_counts_per_range]
            maritals_names = [MaritalStatus.objects.get(id=item['marital_status']).maritals_name for item in gender_counts_per_marital]
            educ_names = [EducAttain.objects.get(id=item['educational_attainment']).educ_name for item in gender_counts_per_attainment]


            female_count = [item['count'] for item in gender_counts]


            marital_status_counts = [item['gender_count'] for item in gender_counts_per_marital]

            social_data_counts = [item['social_count'] for item in gender_counts_per_social]

            range_data_counts = [item['range_count'] for item in gender_counts_per_range]

            community_count = [item['type_count'] for item in gender_counts_per_type]
            
            educ_data_counts = [item['educ_count'] for item in gender_counts_per_attainment]

            prob_data = [item['problem'] for item in gender_counts_per_prob]
            prob_data_counts = [item['prob_counts'] for item in gender_counts_per_prob]

            means_data = [item['means_of_income'] for item in gender_counts_per_means]
            means_data_counts = [item['means_count'] for item in gender_counts_per_means]
            
            
            data = {
                'gndr_names':gndr_names,
                'female_count':female_count,
                'type_names':type_names,
                'community_count':community_count,
                'social_names':social_names,
                'maritals_names':maritals_names,
                'marital_status_counts':marital_status_counts,
                'social_data_counts':social_data_counts,
                'ranges_names':ranges_names,
                'range_data_counts':range_data_counts,
                'means_data':means_data,
                'means_data_counts':means_data_counts,
                'educ_names':educ_names,
                'educ_data_counts':educ_data_counts,
                'prob_data':prob_data,
                'prob_data_counts':prob_data_counts
            }

        elif filter1 == 'Type' and filter_id == filter_id:
                gender_counts_per_marital = PersonalInfo.objects.values('marital_status').annotate(
                        gender_count=Count('type', filter=Q(type=filter_id)),
                    )
                
                gender_counts_per_gender= PersonalInfo.objects.values('sex_gender').annotate(
                        gender_count=Count('type', filter=Q(type=filter_id)),
                    )

                gender_counts_per_social= PersonalInfo.objects.values('estimated_income').annotate(
                        social_count=Count('type', filter=Q(type=filter_id)),
                    )
                gender_counts_per_range= PersonalInfo.objects.values('range').annotate(
                        range_count=Count('type', filter=Q(type=filter_id)),
                    )
                gender_counts_per_type= PersonalInfo.objects.values('type').annotate(
                        type_count=Count('type', filter=Q(type=filter_id)),
                    )
                gender_counts_per_attainment= PersonalInfo.objects.values('educational_attainment').annotate(
                        educ_count=Count('type', filter=Q(type=filter_id)),
                    )
                gender_counts_per_means= PersonalInfo.objects.values('means_of_income').annotate(
                        means_count=Count('type', filter=Q(type=filter_id)),
                    )
                gender_counts_per_prob= PersonalInfo.objects.values('problem').annotate(
                        prob_counts=Count('type', filter=Q(type=filter_id)),
                    )
                
                gndr_names = [Genders.objects.get(id=item['sex_gender']).gndr_name for item in gender_counts_per_gender]
                type_names = [Types.objects.get(id=item['type']).type_name for item in gender_counts_per_type]
                social_names = [EstIncome.objects.get(id=item['estimated_income']).est_income_name for item in gender_counts_per_social]
                ranges_names = [Ranges.objects.get(id=item['range']).ranges_name for item in gender_counts_per_range]
                maritals_names = [MaritalStatus.objects.get(id=item['marital_status']).maritals_name for item in gender_counts_per_marital]
                educ_names = [EducAttain.objects.get(id=item['educational_attainment']).educ_name for item in gender_counts_per_attainment]

                female_count = [item['gender_count'] for item in gender_counts_per_gender]


                marital_status_counts = [item['gender_count'] for item in gender_counts_per_marital]

                social_data_counts = [item['social_count'] for item in gender_counts_per_social]

                range_data_counts = [item['range_count'] for item in gender_counts_per_range]

                community_count = [item['type_count'] for item in gender_counts_per_type]
                
                educ_data_counts = [item['educ_count'] for item in gender_counts_per_attainment]

                prob_data = [item['problem'] for item in gender_counts_per_prob]
                prob_data_counts = [item['prob_counts'] for item in gender_counts_per_prob]

                means_data = [item['means_of_income'] for item in gender_counts_per_means]
                means_data_counts = [item['means_count'] for item in gender_counts_per_means]
                 

                data = {
                'gndr_names':gndr_names,
                'female_count':female_count,
                'type_names':type_names,
                'community_count':community_count,
                'social_names':social_names,
                'maritals_names':maritals_names,
                'marital_status_counts':marital_status_counts,
                'social_data_counts':social_data_counts,
                'ranges_names':ranges_names,
                'range_data_counts':range_data_counts,
                'means_data':means_data,
                'means_data_counts':means_data_counts,
                'educ_names':educ_names,
                'educ_data_counts':educ_data_counts,
                'prob_data':prob_data,
                'prob_data_counts':prob_data_counts
            }

        elif filter1 == 'Age Range' and filter_id == filter_id:
                gender_counts_per_marital = PersonalInfo.objects.values('marital_status').annotate(
                        gender_count=Count('range', filter=Q(range=filter_id)),
                    )
                
                gender_counts_per_gender= PersonalInfo.objects.values('sex_gender').annotate(
                        gender_count=Count('range', filter=Q(range=filter_id)),
                    )

                gender_counts_per_social= PersonalInfo.objects.values('estimated_income').annotate(
                        social_count=Count('range', filter=Q(range=filter_id)),
                    )
                gender_counts_per_range= PersonalInfo.objects.values('range').annotate(
                        range_count=Count('range', filter=Q(range=filter_id)),
                    )
                gender_counts_per_type= PersonalInfo.objects.values('type').annotate(
                        type_count=Count('range', filter=Q(range=filter_id)),
                    )
                gender_counts_per_attainment= PersonalInfo.objects.values('educational_attainment').annotate(
                        educ_count=Count('range', filter=Q(range=filter_id)),
                    )
                gender_counts_per_means= PersonalInfo.objects.values('means_of_income').annotate(
                        means_count=Count('range', filter=Q(range=filter_id)),
                    )
                gender_counts_per_prob= PersonalInfo.objects.values('problem').annotate(
                        prob_counts=Count('range', filter=Q(range=filter_id)),
                    )
                
                gndr_names = [Genders.objects.get(id=item['sex_gender']).gndr_name for item in gender_counts_per_gender]
                type_names = [Types.objects.get(id=item['type']).type_name for item in gender_counts_per_type]
                social_names = [EstIncome.objects.get(id=item['estimated_income']).est_income_name for item in gender_counts_per_social]
                ranges_names = [Ranges.objects.get(id=item['range']).ranges_name for item in gender_counts_per_range]
                maritals_names = [MaritalStatus.objects.get(id=item['marital_status']).maritals_name for item in gender_counts_per_marital]
                educ_names = [EducAttain.objects.get(id=item['educational_attainment']).educ_name for item in gender_counts_per_attainment]

                female_count = [item['gender_count'] for item in gender_counts_per_gender]


                marital_status_counts = [item['gender_count'] for item in gender_counts_per_marital]

                social_data_counts = [item['social_count'] for item in gender_counts_per_social]

                range_data_counts = [item['range_count'] for item in gender_counts_per_range]

                community_count = [item['type_count'] for item in gender_counts_per_type]
                
                educ_data_counts = [item['educ_count'] for item in gender_counts_per_attainment]

                prob_data = [item['problem'] for item in gender_counts_per_prob]
                prob_data_counts = [item['prob_counts'] for item in gender_counts_per_prob]

                means_data = [item['means_of_income'] for item in gender_counts_per_means]
                means_data_counts = [item['means_count'] for item in gender_counts_per_means]
                

                data = {
                    'gndr_names':gndr_names,
                    'female_count':female_count,
                    'type_names':type_names,
                    'community_count':community_count,
                    'social_names':social_names,
                    'maritals_names':maritals_names,
                    'marital_status_counts':marital_status_counts,
                    'social_data_counts':social_data_counts,
                    'ranges_names':ranges_names,
                    'range_data_counts':range_data_counts,
                    'means_data':means_data,
                    'means_data_counts':means_data_counts,
                    'educ_names':educ_names,
                    'educ_data_counts':educ_data_counts,
                    'prob_data':prob_data,
                    'prob_data_counts':prob_data_counts
                }

        elif filter1 == 'Social Status' and filter_id == filter_id:
                gender_counts_per_marital = PersonalInfo.objects.values('marital_status').annotate(
                        gender_count=Count('estimated_income', filter=Q(estimated_income=filter_id)),
                    )
                gender_counts_per_gender= PersonalInfo.objects.values('sex_gender').annotate(
                        gender_count=Count('estimated_income', filter=Q(estimated_income=filter_id)),
                    )

                gender_counts_per_social= PersonalInfo.objects.values('estimated_income').annotate(
                        social_count=Count('estimated_income', filter=Q(estimated_income=filter_id)),
                    )
                gender_counts_per_range= PersonalInfo.objects.values('range').annotate(
                        range_count=Count('estimated_income', filter=Q(estimated_income=filter_id)),
                    )
                gender_counts_per_type= PersonalInfo.objects.values('type').annotate(
                        type_count=Count('estimated_income', filter=Q(estimated_income=filter_id)),
                    )
                gender_counts_per_attainment= PersonalInfo.objects.values('educational_attainment').annotate(
                        educ_count=Count('estimated_income', filter=Q(estimated_income=filter_id)),
                    )
                gender_counts_per_means= PersonalInfo.objects.values('means_of_income').annotate(
                        means_count=Count('estimated_income', filter=Q(estimated_income=filter_id)),
                    )
                gender_counts_per_prob= PersonalInfo.objects.values('problem').annotate(
                        prob_counts=Count('estimated_income', filter=Q(estimated_income=filter_id)),
                    )
                
                gndr_names = [Genders.objects.get(id=item['sex_gender']).gndr_name for item in gender_counts_per_gender]
                type_names = [Types.objects.get(id=item['type']).type_name for item in gender_counts_per_type]
                social_names = [EstIncome.objects.get(id=item['estimated_income']).est_income_name for item in gender_counts_per_social]
                ranges_names = [Ranges.objects.get(id=item['range']).ranges_name for item in gender_counts_per_range]
                maritals_names = [MaritalStatus.objects.get(id=item['marital_status']).maritals_name for item in gender_counts_per_marital]
                educ_names = [EducAttain.objects.get(id=item['educational_attainment']).educ_name for item in gender_counts_per_attainment]
                
                female_count = [item['gender_count'] for item in gender_counts_per_gender]

                marital_status_counts = [item['gender_count'] for item in gender_counts_per_marital]

                social_data_counts = [item['social_count'] for item in gender_counts_per_social]

                range_data_counts = [item['range_count'] for item in gender_counts_per_range]

                community_count = [item['type_count'] for item in gender_counts_per_type]
                
                educ_data_counts = [item['educ_count'] for item in gender_counts_per_attainment]

                prob_data = [item['problem'] for item in gender_counts_per_prob]
                prob_data_counts = [item['prob_counts'] for item in gender_counts_per_prob]

                means_data = [item['means_of_income'] for item in gender_counts_per_means]
                means_data_counts = [item['means_count'] for item in gender_counts_per_means]
                

                data = {
                    'gndr_names':gndr_names,
                    'female_count':female_count,
                    'type_names':type_names,
                    'community_count':community_count,
                    'social_names':social_names,
                    'maritals_names':maritals_names,
                    'marital_status_counts':marital_status_counts,
                    'social_data_counts':social_data_counts,
                    'ranges_names':ranges_names,
                    'range_data_counts':range_data_counts,
                    'means_data':means_data,
                    'means_data_counts':means_data_counts,
                    'educ_names':educ_names,
                    'educ_data_counts':educ_data_counts,
                    'prob_data':prob_data,
                    'prob_data_counts':prob_data_counts
                }
        
        elif filter1 == 'Marital Status' and filter_id == filter_id:
                gender_counts_per_marital = PersonalInfo.objects.values('marital_status').annotate(
                        gender_count=Count('marital_status', filter=Q(marital_status=filter_id)),
                    )
                
                gender_counts_per_gender= PersonalInfo.objects.values('sex_gender').annotate(
                        gender_count=Count('marital_status', filter=Q(marital_status=filter_id)),
                    )

                gender_counts_per_social= PersonalInfo.objects.values('estimated_income').annotate(
                        social_count=Count('marital_status', filter=Q(marital_status=filter_id)),
                    )
                gender_counts_per_range= PersonalInfo.objects.values('range').annotate(
                        range_count=Count('marital_status', filter=Q(marital_status=filter_id)),
                    )
                gender_counts_per_type= PersonalInfo.objects.values('type').annotate(
                        type_count=Count('marital_status', filter=Q(marital_status=filter_id)),
                    )
                gender_counts_per_attainment= PersonalInfo.objects.values('educational_attainment').annotate(
                        educ_count=Count('marital_status', filter=Q(marital_status=filter_id)),
                    )
                gender_counts_per_means= PersonalInfo.objects.values('means_of_income').annotate(
                        means_count=Count('marital_status', filter=Q(marital_status=filter_id)),
                    )
                gender_counts_per_prob= PersonalInfo.objects.values('problem').annotate(
                        prob_counts=Count('marital_status', filter=Q(marital_status=filter_id)),
                    )
                
                gndr_names = [Genders.objects.get(id=item['sex_gender']).gndr_name for item in gender_counts_per_gender]
                type_names = [Types.objects.get(id=item['type']).type_name for item in gender_counts_per_type]
                social_names = [EstIncome.objects.get(id=item['estimated_income']).est_income_name for item in gender_counts_per_social]
                ranges_names = [Ranges.objects.get(id=item['range']).ranges_name for item in gender_counts_per_range]
                maritals_names = [MaritalStatus.objects.get(id=item['marital_status']).maritals_name for item in gender_counts_per_marital]
                educ_names = [EducAttain.objects.get(id=item['educational_attainment']).educ_name for item in gender_counts_per_attainment]
                
                female_count = [item['gender_count'] for item in gender_counts_per_gender]

                marital_status_counts = [item['gender_count'] for item in gender_counts_per_marital]

                social_data_counts = [item['social_count'] for item in gender_counts_per_social]

                range_data_counts = [item['range_count'] for item in gender_counts_per_range]

                community_count = [item['type_count'] for item in gender_counts_per_type]
                
                educ_data_counts = [item['educ_count'] for item in gender_counts_per_attainment]

                prob_data = [item['problem'] for item in gender_counts_per_prob]
                prob_data_counts = [item['prob_counts'] for item in gender_counts_per_prob]

                means_data = [item['means_of_income'] for item in gender_counts_per_means]
                means_data_counts = [item['means_count'] for item in gender_counts_per_means]
                

                data = {
                    'gndr_names':gndr_names,
                    'female_count':female_count,
                    'type_names':type_names,
                    'community_count':community_count,
                    'social_names':social_names,
                    'maritals_names':maritals_names,
                    'marital_status_counts':marital_status_counts,
                    'social_data_counts':social_data_counts,
                    'ranges_names':ranges_names,
                    'range_data_counts':range_data_counts,
                    'means_data':means_data,
                    'means_data_counts':means_data_counts,
                    'educ_names':educ_names,
                    'educ_data_counts':educ_data_counts,
                    'prob_data':prob_data,
                    'prob_data_counts':prob_data_counts
                }
        
        elif filter1 == 'Educational Attainment' and filter_id == filter_id:
                gender_counts_per_marital = PersonalInfo.objects.values('marital_status').annotate(
                        gender_count=Count('educational_attainment', filter=Q(educational_attainment=filter_id)),
                    )
                
                gender_counts_per_gender= PersonalInfo.objects.values('sex_gender').annotate(
                        gender_count=Count('educational_attainment', filter=Q(educational_attainment=filter_id)),
                    )

                gender_counts_per_social= PersonalInfo.objects.values('estimated_income').annotate(
                        social_count=Count('educational_attainment', filter=Q(educational_attainment=filter_id)),
                    )
                gender_counts_per_range= PersonalInfo.objects.values('range').annotate(
                        range_count=Count('educational_attainment', filter=Q(educational_attainment=filter_id)),
                    )
                gender_counts_per_type= PersonalInfo.objects.values('type').annotate(
                        type_count=Count('educational_attainment', filter=Q(educational_attainment=filter_id)),
                    )
                gender_counts_per_attainment= PersonalInfo.objects.values('educational_attainment').annotate(
                        educ_count=Count('educational_attainment', filter=Q(educational_attainment=filter_id)),
                    )
                gender_counts_per_means= PersonalInfo.objects.values('means_of_income').annotate(
                        means_count=Count('educational_attainment', filter=Q(educational_attainment=filter_id)),
                    )
                gender_counts_per_prob= PersonalInfo.objects.values('problem').annotate(
                        prob_counts=Count('educational_attainment', filter=Q(educational_attainment=filter_id)),
                    )
                
                gndr_names = [Genders.objects.get(id=item['sex_gender']).gndr_name for item in gender_counts_per_gender]
                type_names = [Types.objects.get(id=item['type']).type_name for item in gender_counts_per_type]
                social_names = [EstIncome.objects.get(id=item['estimated_income']).est_income_name for item in gender_counts_per_social]
                ranges_names = [Ranges.objects.get(id=item['range']).ranges_name for item in gender_counts_per_range]
                maritals_names = [MaritalStatus.objects.get(id=item['marital_status']).maritals_name for item in gender_counts_per_marital]
                educ_names = [EducAttain.objects.get(id=item['educational_attainment']).educ_name for item in gender_counts_per_attainment]


                female_count = [item['gender_count'] for item in gender_counts_per_gender]


                marital_status_counts = [item['gender_count'] for item in gender_counts_per_marital]

                social_data_counts = [item['social_count'] for item in gender_counts_per_social]

                range_data_counts = [item['range_count'] for item in gender_counts_per_range]

                community_count = [item['type_count'] for item in gender_counts_per_type]
                
                educ_data_counts = [item['educ_count'] for item in gender_counts_per_attainment]

                prob_data = [item['problem'] for item in gender_counts_per_prob]
                prob_data_counts = [item['prob_counts'] for item in gender_counts_per_prob]

                means_data = [item['means_of_income'] for item in gender_counts_per_means]
                means_data_counts = [item['means_count'] for item in gender_counts_per_means]
                

                data = {
                    'gndr_names':gndr_names,
                    'female_count':female_count,
                    'type_names':type_names,
                    'community_count':community_count,
                    'social_names':social_names,
                    'maritals_names':maritals_names,
                    'marital_status_counts':marital_status_counts,
                    'social_data_counts':social_data_counts,
                    'ranges_names':ranges_names,
                    'range_data_counts':range_data_counts,
                    'means_data':means_data,
                    'means_data_counts':means_data_counts,
                    'educ_names':educ_names,
                    'educ_data_counts':educ_data_counts,
                    'prob_data':prob_data,
                    'prob_data_counts':prob_data_counts
                }

        print(data)
        return JsonResponse(data)
    return JsonResponse({'error': 'Invalid request'})

def search(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        searched = data.get('search')
        folders = PersonalInfo.objects.filter(problem__icontains=searched).values('problem').annotate(count=Count('id')).order_by('-count')

        folder_names = [{'problem': folder['problem'], 'count': folder['count']} for folder in folders]
        
        data = {
            'folder_names': folder_names,
        }
        return JsonResponse(data)
    
from django.db.models import Q

def search_file(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        searched = data.get('search')
        searched2 = data.get('search2')

        files = PersonalInfo.objects.filter(
            (Q(full_name__icontains=searched) |
             Q(type__type_name__icontains=searched) |  
             Q(sex_gender__gndr_name__icontains=searched)) & 
            (Q(problem__icontains=searched2))
        )

        results = []
        for file_obj in files:
            document_value = file_obj.document.url if file_obj.document else None 
            result = {
                'full_name': file_obj.full_name,
                'problem': file_obj.problem,
                'document': document_value,
            }
            results.append(result)

        data = {
            'results': results,
        }

        return JsonResponse(data)
    

def get(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        prob = data.get('problems')
        files = PersonalInfo.objects.filter(problem=prob)

        results = []
        for file_obj in files:
            document_value = file_obj.document.url if file_obj.document else None 
            result = {
                'full_name': file_obj.full_name,
                'problem': file_obj.problem,
                'document': document_value,
            }
            results.append(result)

        data = {
            'results': results,
        }

        print(data)

        return JsonResponse(data)
def addGen(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        gender = data.get('gndr')
        
        if not Genders.objects.filter(gndr_name=gender).exists():
            Genders.objects.create(gndr_name=gender)

        print(gender)
        print("added!")

        getGenders = Genders.objects.all()

        serialized_genders = serialize('json', getGenders)

        return JsonResponse({'getGenders': serialized_genders})
    
def fetchGen(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        gender = data.get('gndr')
        
        gender = Genders.objects.get(id=gender)

        gndr_name = gender.gndr_name 
        return JsonResponse({'gndr_name': gndr_name})

    
def delGen(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        gender = data.get('gndr')
        
        gender = Genders.objects.filter(id=gender)
        
        if gender is not None:
            gender.delete()

            getGenders = Genders.objects.all()

            serialized_genders = serialize('json', getGenders)

            return JsonResponse({'getGenders': serialized_genders})
        else:
            return JsonResponse({'error': 'Gender not found'}, status=404)


def addMar(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        marital = data.get('mrtl')
        
        if not MaritalStatus.objects.filter(maritals_name=marital).exists():
            MaritalStatus.objects.create(maritals_name=marital)

        print(marital)
        print("added!")

        getMaritals = MaritalStatus.objects.all()

        serialized_maritals = serialize('json', getMaritals)

        return JsonResponse({'getMaritals': serialized_maritals})
    
def fetchMar(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        marital = data.get('mrtl')
        
        marital = MaritalStatus.objects.get(id=marital)

        mrtl_name = marital.maritals_name 
        return JsonResponse({'mrtl_name': mrtl_name})
    
def delMar(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        marital = data.get('mrtl')
        
        maritals = MaritalStatus.objects.filter(id=marital)
        
        if maritals is not None:
            maritals.delete()

            getMaritals = MaritalStatus.objects.all()

            serialized_maritals = serialize('json', getMaritals)

            return JsonResponse({'getMaritals': serialized_maritals})
        else:
            return JsonResponse({'error': 'Gender not found'}, status=404)

def addRange(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        rng = data.get('rng')
        rngdet = data.get('rngDet')
        rngdet1 = data.get('rngDet1')

        toRagne = rngdet + ' to ' + rngdet1

        if not Ranges.objects.filter(ranges_name=rng, ranges_details=toRagne).exists():
            Ranges.objects.create(ranges_name=rng,range1=rngdet, range2=rngdet1, ranges_details=toRagne)
            
        print("added!")

        getRanges = Ranges.objects.all()

        serialized_ranges = serialize('json', getRanges)

        return JsonResponse({'getRanges': serialized_ranges})
    
def editRanges(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        ranges = data.get('rngs')
        ranges = Ranges.objects.get(id=ranges)

        range_names = ranges.ranges_name
        range1s = ranges.range1
        range2s = ranges.range2
        data = {
                'range_names': range_names,
                'range1s':range1s,
                'range2s':range2s
        }
        return JsonResponse(data)

def delRange(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        range = data.get('rng')
        
        ranges = Ranges.objects.filter(id=range)
        
        if ranges is not None:
            ranges.delete()

            getRanges = Ranges.objects.all()

            serialized_Ranges = serialize('json', getRanges)

            return JsonResponse({'getRanges': serialized_Ranges})
        else:
            return JsonResponse({'error': 'Gender not found'}, status=404)
        
        
def addEduc(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        educ = data.get('educ')
        
        if not EducAttain.objects.filter(educ_name=educ).exists():
            EducAttain.objects.create(educ_name=educ)

        print(educ)
        print("added!")

        getEduc = EducAttain.objects.all()

        serialized_educs = serialize('json', getEduc)

        return JsonResponse({'getEduc': serialized_educs})
    
def fetchEduc(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        educ = data.get('educId')
        
        educ = EducAttain.objects.get(id=educ)

        educ_name = educ.educ_name 
        return JsonResponse({'educ_name': educ_name})
    
def delEduc(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        educ = data.get('educ')
        
        educs = EducAttain.objects.filter(id=educ)
        
        if educs is not None:
            educs.delete()

            getEduc = EducAttain.objects.all()

            serialized_educ = serialize('json', getEduc)

            return JsonResponse({'getEduc': serialized_educ})
        else:
            return JsonResponse({'error': 'Gender not found'}, status=404)
        
def addSocial(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        scl = data.get('est')
        soc = data.get('soc')
        
        if not EstIncome.objects.filter(est_income_name=soc, est_income_details=scl).exists():
            EstIncome.objects.create(est_income_name=soc, est_income_details=scl)
            
        print("added!")

        getSocials = EstIncome.objects.all()

        serialized_Socials = serialize('json', getSocials)

        return JsonResponse({'getSocials': serialized_Socials})
    
def fetchSocial(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        socId = data.get('socId')
        
        social = EstIncome.objects.get(id=socId)

        social_name = social.est_income_name 
        social_det = social.est_income_details 
        return JsonResponse({'social_name': social_name, 'social_det':social_det})
    
def delSocial(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        social = data.get('scl')
        
        socials = EstIncome.objects.filter(id=social)
        
        if socials is not None:
            socials.delete()

            getSocials = EstIncome.objects.all()

            serialized_Socials = serialize('json', getSocials)

            return JsonResponse({'getSocials': serialized_Socials})
        else:
            return JsonResponse({'error': 'Gender not found'}, status=404)
        
def fetchRequest(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        reqID = data.get('reqID')
        
        request = Requests.objects.get(id=reqID)

        username = request.username 
        email = request.email
        password = request.password 

        if User.objects.filter(username=username).exists() or User.objects.filter(email=email).exists():
            return redirect('admin')
        else:
            my_user=User.objects.create_user(username,email,password)
            my_user.save()

        request.delete()
        return JsonResponse({'status': 'ok'})
    else:
        return JsonResponse({'error': 'Method not allowed'}, status=405)

def deleteRequest(request):
    if request.method == 'POST':
        data = json.loads(request.body)  
        reqID = data.get('reqID')
        
        request = Requests.objects.get(id=reqID)
        request.delete()
        return JsonResponse({'status': 'ok'})
    else:
        return JsonResponse({'error': 'Method not allowed'}, status=405)


def logoutview(request):
    logout(request) 
    return redirect ('landing')  


def AboutPage(request):
    return render(request, 'aboutUs.html') 


def ContactPage(request):
    return render(request, 'contact.html')

def save_document(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        content = data.get('content')
        filename = 'document'
        personal_info_id = data.get('personalInfoId')
        print(filename)
        print(personal_info_id)


        if content and filename and personal_info_id:
            try:
                personal_info = PersonalInfo.objects.get(pk=personal_info_id)

                unique_filename = slugify(filename)
                
                personal_info.document.save(unique_filename, ContentFile(content.encode('utf-8')))  

                return HttpResponse(status=200)
            except PersonalInfo.DoesNotExist:
                return JsonResponse({'error': 'PersonalInfo object not found.'}, status=404)
            except Exception as e:
                return JsonResponse({'error': str(e)}, status=500)
        else:
            return JsonResponse({'error': 'Invalid data.'}, status=400)
    return JsonResponse({'error': 'Bad request.'}, status=400)